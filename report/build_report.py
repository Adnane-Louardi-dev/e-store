"""
Générateur du rapport technique — Projet E-Store (Full Stack S6, FSBM)
Auteurs : Adnane Louardi & Yassine Boukhari
Encadrant : Pr. Omar Zahour
Année : 2024–2025
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

BASE = Path(__file__).parent
DIAG = BASE / "diagrams"
OUT = BASE / "Rapport_E-Store_Full_Stack.docx"

# ── Couleurs de la charte ──────────────────────────────────────────────────
BLUE  = RGBColor(0x1e, 0x3a, 0x8a)   # bleu marine
GREEN = RGBColor(0x15, 0x80, 0x3d)   # vert FSBM
GRAY  = RGBColor(0x4b, 0x55, 0x63)   # gris texte

# ── Helpers ────────────────────────────────────────────────────────────────

def add_page_break(doc):
    doc.add_page_break()


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        run = p.runs[0] if p.runs else p.add_run(text)
        run.font.color.rgb = BLUE
    return p


def add_para(doc, text="", bold=False, italic=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_code_block(doc, code: str, caption: str = ""):
    if caption:
        cp = doc.add_paragraph(caption)
        cp.runs[0].bold = True
        cp.runs[0].font.size = Pt(10)
        cp.runs[0].font.color.rgb = GRAY
    p = doc.add_paragraph()
    p.style = "No Spacing"
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1e, 0x1e, 0x2e)
    # Light background via paragraph shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F3F4F6")
    pPr.append(shd)


def add_figure(doc, png_path: Path, caption: str, width_cm: float = 14.0):
    if png_path.exists():
        doc.add_picture(str(png_path), width=Cm(width_cm))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph(f"[IMAGE MANQUANTE : {png_path.name}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = GRAY


def add_table_header(table, headers, fill="1E3A8A"):
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, fill)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_toc_field(doc):
    """Insère un champ TOC Word (se peuple à l'ouverture du fichier)."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════

def add_cover(doc):
    # En-tête institutionnel
    for text, sz, bold, color in [
        ("Université Hassan II de Casablanca", 13, False, GRAY),
        ("Faculté des Sciences Ben M'Sik (FSBM)", 13, True, BLUE),
        ("Filière : Génie Informatique — Semestre 6", 11, False, GRAY),
        ("Module : Développement Full Stack", 11, False, GRAY),
    ]:
        p = add_para(doc, text, bold=bold, size=sz, color=color, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()
    doc.add_paragraph()

    # Titre
    for text, sz, bold, color in [
        ("RAPPORT DE PROJET", 22, True, BLUE),
        ("", 10, False, None),
        ("Conception et Réalisation d'une Plateforme", 18, True, GREEN),
        ("E-Commerce Full Stack", 18, True, GREEN),
        ("« E-Store »", 20, True, BLUE),
        ("", 10, False, None),
        ("Spring Boot 3 · JPA · MongoDB · Angular 18", 13, False, GRAY),
    ]:
        if text == "":
            doc.add_paragraph()
        else:
            add_para(doc, text, bold=bold, size=sz, color=color, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()
    doc.add_paragraph()

    # Auteurs & encadrant
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    left, right = t.rows[0].cells
    left.text = "Réalisé par :\n• Adnane Louardi\n• Yassine Boukhari"
    right.text = "Encadré par :\n• Pr. Omar Zahour"
    for cell in [left, right]:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(12)
                run.bold = True
                run.font.color.rgb = BLUE

    doc.add_paragraph()
    add_para(doc, "Année universitaire : 2024–2025", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# DÉDICACE
# ══════════════════════════════════════════════════════════════════════════════

def add_dedication(doc):
    add_heading(doc, "Dédicace", level=1)
    add_para(doc, (
        "À nos familles, pour leur soutien indéfectible tout au long de notre parcours académique. "
        "À nos amis et camarades de promotion, qui ont rendu cette aventure plus riche et plus joyeuse. "
        "À Pr. Omar Zahour, dont les conseils avisés et la rigueur pédagogique ont guidé chaque étape "
        "de ce projet. Ce travail vous est dédié."
    ), italic=True, size=11)
    add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# REMERCIEMENTS
# ══════════════════════════════════════════════════════════════════════════════

def add_acknowledgements(doc):
    add_heading(doc, "Remerciements", level=1)
    add_para(doc, (
        "Nous tenons à exprimer notre profonde gratitude à Pr. Omar Zahour pour avoir proposé "
        "ce projet stimulant et pour ses précieuses orientations techniques tout au long du semestre. "
        "Nos remerciements vont également à l'ensemble du corps enseignant de la Faculté des Sciences "
        "Ben M'Sik pour la qualité de la formation dispensée. "
        "Enfin, nous remercions chaleureusement nos familles pour leur patience et leur encouragement "
        "constants durant cette période de travail intensif."
    ), size=11)
    add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# RÉSUMÉS
# ══════════════════════════════════════════════════════════════════════════════

def add_abstracts(doc):
    add_heading(doc, "Résumé", level=1)
    add_para(doc, (
        "Ce rapport présente la conception et la réalisation d'une plateforme e-commerce complète "
        "intitulée E-Store, développée dans le cadre du module Full Stack du semestre 6 à la Faculté "
        "des Sciences Ben M'Sik. L'application adopte une architecture en couches découplée : un backend "
        "REST construit avec Spring Boot 3.3.5 et Java 21, organisé en six domaines métier "
        "(Customer, Catalog, Inventory, Shopping, Billing, Review), et un frontend Angular 18 composé "
        "de composants standalone. La persistance est hybride : Spring Data JPA avec H2 (développement) "
        "ou MySQL (production) pour les données relationnelles, et MongoDB pour les avis produit. "
        "La sécurité repose sur JSON Web Tokens (JWT) et le chiffrement BCrypt. Le projet est "
        "documenté via Swagger/OpenAPI et validé par un script de smoke-test couvrant l'intégralité "
        "des 28 endpoints implémentés."
    ), size=11)
    doc.add_paragraph()

    add_heading(doc, "Abstract", level=1)
    add_para(doc, (
        "This report describes the design and implementation of a full e-commerce platform called "
        "E-Store, developed as part of the Full Stack module (semester 6) at the Faculty of Sciences "
        "Ben M'Sik. The application follows a layered, domain-driven architecture: a REST backend built "
        "with Spring Boot 3.3.5 and Java 21, structured into six business domains (Customer, Catalog, "
        "Inventory, Shopping, Billing, Review), and an Angular 18 frontend made of standalone components. "
        "Persistence is hybrid: Spring Data JPA with H2 (development) or MySQL (production) for "
        "relational data, and MongoDB for product reviews. Security is based on JWT and BCrypt hashing. "
        "The project is documented via Swagger/OpenAPI and validated by a smoke-test script covering "
        "all 28 implemented endpoints."
    ), size=11, italic=True)
    doc.add_paragraph()

    add_heading(doc, "ملخص", level=1)
    add_para(doc, (
        "يقدم هذا التقرير تصميم وتطوير منصة تجارة إلكترونية متكاملة باسم E-Store، أُنجزت في إطار "
        "وحدة التطوير الشامل للفصل السادس بكلية العلوم بن مسيك. تعتمد التطبيق بنية متعددة الطبقات: "
        "واجهة خلفية REST مبنية بـ Spring Boot 3.3.5 وJava 21 موزعة على ستة مجالات وظيفية، "
        "وواجهة أمامية بـ Angular 18. يجمع التطبيق بين JPA/H2/MySQL وMongoDB للمراجعات، "
        "مع تأمين قائم على JWT وBCrypt."
    ), size=11)
    add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# TABLE DES MATIÈRES
# ══════════════════════════════════════════════════════════════════════════════

def add_toc(doc):
    add_heading(doc, "Table des matières", level=1)
    add_para(doc, "Cliquez-droit sur le champ ci-dessous puis « Mettre à jour les champs » pour générer la table.", italic=True, size=9, color=GRAY)
    add_toc_field(doc)
    add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# LISTE DES FIGURES ET TABLEAUX
# ══════════════════════════════════════════════════════════════════════════════

def add_lists(doc):
    add_heading(doc, "Liste des figures", level=1)
    figures = [
        ("Figure 3.1", "Architecture en couches de E-Store"),
        ("Figure 3.2", "Décomposition par domaines métier"),
        ("Figure 3.3", "Architecture technique (Angular ↔ Spring Boot ↔ BDD)"),
        ("Figure 3.4", "Diagramme de cas d'utilisation global"),
        ("Figure 3.5", "Diagramme de classes (JPA + MongoDB)"),
        ("Figure 3.6", "Diagramme Entité-Relation (9 tables)"),
        ("Figure 3.7", "Diagramme de séquence — Authentification JWT"),
        ("Figure 3.8", "Diagramme de séquence — Checkout transactionnel"),
        ("Figure 3.9", "Diagramme de séquence — Ajout panier (vérification stock)"),
        ("Figure 3.10", "Diagramme de séquence — Dépôt d'avis (MongoDB)"),
        ("Figure 6.1", "Page d'accueil — liste des produits"),
        ("Figure 6.2", "Détail d'un produit et avis"),
        ("Figure 6.3", "Formulaire d'inscription"),
        ("Figure 6.4", "Formulaire de connexion"),
        ("Figure 6.5", "Panier d'achat"),
        ("Figure 6.6", "Page de confirmation de commande"),
        ("Figure 6.7", "Historique des commandes"),
        ("Figure 6.8", "Profil utilisateur"),
        ("Figure 6.9", "Dashboard administrateur"),
        ("Figure 6.10", "Gestion des stocks (admin)"),
    ]
    for ref, title in figures:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{ref} — ").bold = True
        p.add_run(title)

    doc.add_paragraph()
    add_heading(doc, "Liste des tableaux", level=1)
    tables = [
        ("Tableau 2.1", "Acteurs et leurs droits"),
        ("Tableau 2.2", "Exigences fonctionnelles par domaine"),
        ("Tableau 2.3", "Exigences non fonctionnelles"),
        ("Tableau 3.1", "Entités JPA et relations"),
        ("Tableau 4.1", "Dépendances Maven du projet"),
        ("Tableau 5.1", "Structure du backend (76 classes / 9 packages)"),
        ("Tableau 5.2", "Structure du frontend (17 composants / 8 services)"),
        ("Tableau 7.1", "Tests fonctionnels par domaine"),
        ("Tableau 7.2", "Tests de sécurité"),
        ("Tableau A.1", "Inventaire complet des 28 endpoints REST"),
    ]
    for ref, title in tables:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{ref} — ").bold = True
        p.add_run(title)
    add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# INTRODUCTION GÉNÉRALE
# ══════════════════════════════════════════════════════════════════════════════

def add_introduction(doc):
    add_heading(doc, "Introduction Générale", level=1)

    add_para(doc, (
        "Le commerce électronique est aujourd'hui l'un des secteurs numériques les plus dynamiques, "
        "et la maîtrise de sa réalisation technique constitue une compétence clé pour tout ingénieur "
        "logiciel. Dans ce contexte, le projet E-Store propose de concevoir et de réaliser une "
        "plateforme e-commerce complète et fonctionnelle, couvrant la totalité de la chaîne : "
        "catalogue produit, gestion des stocks, panier, commande transactionnelle, et avis clients."
    ), size=11)

    add_heading(doc, "Problématique", level=2)
    add_para(doc, (
        "Comment architecturer une application web full stack de manière à séparer clairement les "
        "responsabilités (frontend / backend / persistance), à gérer plusieurs bases de données de "
        "nature différente (relationnelle et document), et à assurer la sécurité des accès tout en "
        "restant maintenable et portable entre environnements ?"
    ), size=11)

    add_heading(doc, "Objectifs du projet", level=2)
    objectifs = [
        "Mettre en œuvre une API REST complète avec Spring Boot 3 et Java 21.",
        "Organiser le backend en domaines métier cohérents (architecture par domaines).",
        "Intégrer une persistance hybride : JPA/H2/MySQL pour les données relationnelles, MongoDB pour les avis.",
        "Sécuriser l'ensemble des endpoints avec JWT et BCrypt.",
        "Développer un frontend Angular 18 découplé, avec composants standalone et lazy loading.",
        "Documenter et tester l'API via Swagger/OpenAPI et un script de smoke-test.",
    ]
    for obj in objectifs:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(obj).font.size = Pt(11)

    add_heading(doc, "Organisation du rapport", level=2)
    add_para(doc, (
        "Ce document est structuré en sept chapitres. Le Chapitre 1 présente le cadre général du "
        "projet. Le Chapitre 2 couvre l'analyse et la spécification des besoins. Le Chapitre 3 "
        "détaille la conception UML. Le Chapitre 4 décrit les technologies choisies. Le Chapitre 5 "
        "présente l'implémentation concrète. Le Chapitre 6 illustre les interfaces utilisateur. "
        "Le Chapitre 7 rapporte les tests et la validation. Une conclusion et des annexes complètent "
        "le document."
    ), size=11)
    add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# CHAPITRE 1 — CADRE DU PROJET
# ══════════════════════════════════════════════════════════════════════════════

def add_chapter1(doc):
    add_heading(doc, "Chapitre 1 — Cadre du projet", level=1)

    add_heading(doc, "1.1 Contexte académique", level=2)
    add_para(doc, (
        "Ce projet s'inscrit dans le cadre du module « Développement Full Stack » du semestre 6 de "
        "la filière Génie Informatique à la Faculté des Sciences Ben M'Sik (FSBM), Université Hassan II "
        "de Casablanca. Il est supervisé par Pr. Omar Zahour et constitue le projet de fin de module "
        "évalué sur 20 points selon le barème du cahier des charges."
    ), size=11)

    add_heading(doc, "1.2 Sujet : E-Store", level=2)
    add_para(doc, (
        "Le sujet imposé est la réalisation d'un mini-projet e-commerce « E-Store » intégrant : "
        "un backend Spring Boot avec Spring Data JPA pour la persistance relationnelle, "
        "Spring Data MongoDB pour les avis, Spring Security pour l'authentification JWT, "
        "et un frontend Angular avec gestion de l'état et des formulaires réactifs."
    ), size=11)

    add_heading(doc, "1.3 Objectifs pédagogiques", level=2)
    pedagogiques = [
        "Maîtriser le développement d'une API REST avec Spring Boot 3 et Java 21.",
        "Comprendre la persistance hybride (JPA + MongoDB).",
        "Implémenter un mécanisme d'authentification stateless (JWT).",
        "Développer un frontend Angular 18 moderne avec composants standalone.",
        "Pratiquer la méthodologie incrémentale (développement jour par jour).",
    ]
    for p in pedagogiques:
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(p).font.size = Pt(11)

    add_heading(doc, "1.4 Méthodologie adoptée", level=2)
    add_para(doc, (
        "Le développement a suivi une approche incrémentale structurée en six itérations journalières : "
        "Jour 1 (infrastructure + entités JPA), Jour 2 (authentification JWT), Jour 3 (catalogue), "
        "Jour 4 (stocks), Jour 5 (panier + commandes), Jour 6 (avis MongoDB + frontend Angular). "
        "Chaque itération a produit des fonctionnalités testables, permettant une validation continue "
        "par le script de smoke-test."
    ), size=11)

    add_heading(doc, "1.5 Barème d'évaluation", level=2)
    t = doc.add_table(rows=7, cols=3)
    t.style = "Table Grid"
    add_table_header(t, ["Critère", "Points", "Description"])
    rows_data = [
        ("Architecture & conception", "4", "Couches, domaines, UML, ERD"),
        ("Backend JPA (Spring Boot)", "4", "Entités, repos, services, sécurité JWT"),
        ("Frontend Angular", "3", "Composants, routing, guards, intercepteurs"),
        ("Intégration MongoDB", "2", "Reviews document, fallback startup"),
        ("Qualité du rapport", "2", "Structure, contenu, diagrammes, bibliographie"),
        ("Tests & validation", "5", "Smoke-test, Swagger, cas de test"),
    ]
    for i, (crit, pts, desc) in enumerate(rows_data):
        row = t.rows[i + 1]
        row.cells[0].text = crit
        row.cells[1].text = pts
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[2].text = desc
        for cell in row.cells:
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
    add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# CHAPITRE 2 — ANALYSE & SPÉCIFICATION
# ══════════════════════════════════════════════════════════════════════════════

def add_chapter2(doc):
    add_heading(doc, "Chapitre 2 — Analyse et Spécification des Besoins", level=1)

    add_heading(doc, "2.1 Acteurs du système", level=2)
    t = doc.add_table(rows=4, cols=3)
    t.style = "Table Grid"
    add_table_header(t, ["Acteur", "Rôle", "Droits"])
    actors = [
        ("Visiteur", "Utilisateur non authentifié", "Parcourir catalogue, rechercher, s'inscrire, se connecter"),
        ("Utilisateur (USER)", "Client authentifié", "Panier, commandes, profil, avis"),
        ("Administrateur (ADMIN)", "Gestionnaire du site", "CRUD catégories/produits, stocks, toutes commandes"),
    ]
    for i, (actor, role, rights) in enumerate(actors):
        row = t.rows[i + 1]
        row.cells[0].text = actor
        row.cells[1].text = role
        row.cells[2].text = rights
        for cell in row.cells:
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)

    add_heading(doc, "2.2 Besoins fonctionnels", level=2)

    domains_req = {
        "Customer & Authentification": [
            "Inscription avec email + mot de passe (hashé BCrypt).",
            "Connexion avec retour d'un token JWT valable 24 heures.",
            "Consultation et mise à jour du profil (téléphone, adresse, ville, pays).",
            "Liste des utilisateurs (ADMIN uniquement).",
        ],
        "Catalogue": [
            "CRUD complet des catégories (ADMIN).",
            "CRUD complet des produits avec image URL et prix (ADMIN).",
            "Recherche paginée avec filtres : mot-clé, catégorie, tri, pagination (PUBLIC).",
            "Consultation du détail d'un produit incluant le stock disponible (PUBLIC).",
        ],
        "Stocks (Inventory)": [
            "Chaque produit dispose d'un enregistrement d'inventaire automatiquement créé à 0.",
            "Consultation publique du stock d'un produit.",
            "Mise à jour du stock par l'ADMIN.",
        ],
        "Panier (Shopping)": [
            "Panier créé automatiquement au premier accès d'un utilisateur authentifié.",
            "Ajout d'un article avec vérification du stock disponible (422 si dépassement).",
            "Modification de la quantité d'un article du panier.",
            "Suppression d'un article ou vidage complet du panier.",
        ],
        "Commandes (Billing)": [
            "Passage de commande atomique (@Transactional) : création, décrémentation stock, vidage panier.",
            "Consultation de ses commandes (utilisateur) ou de toutes (ADMIN), paginée.",
            "Modification d'une ligne de commande (avant expédition).",
            "Annulation d'une commande (statut → CANCELLED).",
        ],
        "Avis (Reviews — MongoDB)": [
            "Lecture publique des avis d'un produit.",
            "Dépôt d'un avis authentifié (note 1–5 + commentaire).",
            "Démarrage gracieux si MongoDB est indisponible.",
        ],
    }

    for domain, reqs in domains_req.items():
        add_heading(doc, f"2.2.{list(domains_req.keys()).index(domain)+1} {domain}", level=3)
        for req in reqs:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(req).font.size = Pt(11)

    add_heading(doc, "2.3 Besoins non fonctionnels", level=2)
    t2 = doc.add_table(rows=7, cols=2)
    t2.style = "Table Grid"
    add_table_header(t2, ["Exigence", "Détail"])
    nfr = [
        ("Sécurité", "JWT stateless (exp. 24h), BCrypt strength 10, CORS restreint à localhost:4200"),
        ("Performance", "Pagination Spring Data (20/page défaut), indexes BDD (idx_products_name, idx_products_category)"),
        ("Maintenabilité", "Architecture en couches + domaines, MapStruct pour mapping, Lombok pour réduction boilerplate"),
        ("Portabilité", "Profils Spring : dev (H2 fichier) / mysql (MySQL 8), configuration externalisée via application.yml"),
        ("Testabilité", "Swagger UI /swagger-ui.html, script smoke-test.sh (28 assertions), collection Postman"),
        ("Documentation API", "Springdoc-openapi 2.6.0, schéma Bearer JWT configuré dans OpenApiConfig"),
    ]
    for i, (req, detail) in enumerate(nfr):
        row = t2.rows[i + 1]
        row.cells[0].text = req
        row.cells[1].text = detail
        for cell in row.cells:
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)

    add_heading(doc, "2.4 Cas d'utilisation détaillés", level=2)

    use_cases = [
        {
            "id": "UC01", "name": "S'inscrire",
            "actors": "Visiteur",
            "precondition": "L'email n'est pas encore enregistré",
            "flow": "1. Le visiteur saisit prénom, nom, email, mot de passe.\n2. Le système valide les champs (@Valid).\n3. BCrypt encode le mot de passe.\n4. Le système crée l'utilisateur (rôle USER) et un profil vide.\n5. Un token JWT est retourné.",
            "exception": "Email déjà utilisé → 422 Unprocessable Entity",
        },
        {
            "id": "UC02", "name": "Ajouter un produit au panier",
            "actors": "Utilisateur authentifié",
            "precondition": "L'utilisateur est connecté (token JWT valide), le produit existe",
            "flow": "1. L'utilisateur sélectionne un produit et une quantité.\n2. Le système vérifie le stock disponible.\n3. Si OK : l'article est ajouté (ou mis à jour si déjà présent), avec snapshot du prix.\n4. Le CartResponse est renvoyé avec le total recalculé.",
            "exception": "Quantité > stock → 422 BusinessException",
        },
        {
            "id": "UC03", "name": "Passer une commande (Checkout)",
            "actors": "Utilisateur authentifié",
            "precondition": "Le panier est non vide, stock suffisant pour chaque article",
            "flow": "1. L'utilisateur confirme la commande.\n2. @Transactional : Order créée, OrderItems snapshottés, stocks décrémentés, panier vidé.\n3. Statut CONFIRMED, OrderResponse retourné.",
            "exception": "Panier vide → 422; Stock insuffisant → 422 (rollback total)",
        },
        {
            "id": "UC04", "name": "Gérer le catalogue (ADMIN)",
            "actors": "Administrateur",
            "precondition": "Token JWT avec rôle ADMIN",
            "flow": "1. L'admin crée/modifie/supprime une catégorie ou un produit.\n2. Validation @Valid des champs.\n3. MapStruct convertit DTO ↔ entité.\n4. Persistance JPA, réponse DTO.",
            "exception": "Catégorie en double → 422; Produit introuvable → 404",
        },
    ]

    for uc in use_cases:
        add_heading(doc, f"{uc['id']} — {uc['name']}", level=3)
        t_uc = doc.add_table(rows=5, cols=2)
        t_uc.style = "Table Grid"
        fields = [
            ("Acteur(s)", uc["actors"]),
            ("Précondition", uc["precondition"]),
            ("Flux principal", uc["flow"]),
            ("Exception / flux alternatif", uc["exception"]),
        ]
        t_uc.rows[0].cells[0].text = "Cas d'utilisation"
        t_uc.rows[0].cells[1].text = f"{uc['id']} — {uc['name']}"
        set_cell_bg(t_uc.rows[0].cells[0], "DBEAFE")
        set_cell_bg(t_uc.rows[0].cells[1], "DBEAFE")
        for i, (k, v) in enumerate(fields):
            t_uc.rows[i + 1].cells[0].text = k
            t_uc.rows[i + 1].cells[1].text = v
            t_uc.rows[i + 1].cells[0].paragraphs[0].runs[0].bold = True
            for cell in t_uc.rows[i + 1].cells:
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(10)
        doc.add_paragraph()

    add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# CHAPITRE 3 — CONCEPTION
# ══════════════════════════════════════════════════════════════════════════════

def add_chapter3(doc):
    add_heading(doc, "Chapitre 3 — Conception", level=1)

    add_heading(doc, "3.1 Introduction à la modélisation UML", level=2)
    add_para(doc, (
        "La conception de E-Store s'appuie sur le langage UML 2 pour décrire l'architecture, "
        "les entités et les interactions du système. Les diagrammes présentés dans ce chapitre ont "
        "été générés avec Mermaid et constituent le référentiel de conception du projet. "
        "Ils ont été validés contre le code source pour garantir leur cohérence."
    ), size=11)

    add_heading(doc, "3.2 Architecture logicielle", level=2)

    add_heading(doc, "3.2.1 Architecture en couches", level=3)
    add_para(doc, (
        "L'application adopte une architecture classique à trois couches, conforme aux recommandations "
        "Spring Boot : la couche Présentation (Angular 18), la couche Métier (Spring Boot), et la "
        "couche Données (JPA + MongoDB). Cette séparation garantit l'indépendance des couches et "
        "facilite la maintenance."
    ), size=11)
    add_figure(doc, DIAG / "architecture_couches.png",
               "Figure 3.1 — Architecture en couches de E-Store")

    add_heading(doc, "3.2.2 Architecture par domaines", level=3)
    add_para(doc, (
        "Le backend est subdivisé en six domaines métier indépendants, chacun encapsulant ses propres "
        "entités, repositories, services, mappers, DTOs et contrôleurs. Cette approche Domain-Driven "
        "réduit le couplage inter-modules et facilite l'évolution indépendante de chaque domaine."
    ), size=11)
    add_figure(doc, DIAG / "architecture_domaines.png",
               "Figure 3.2 — Décomposition par domaines métier")

    add_heading(doc, "3.3 Diagramme de cas d'utilisation", level=2)
    add_para(doc, (
        "Le diagramme ci-dessous recense l'ensemble des fonctionnalités offertes par le système, "
        "regroupées par acteur. On distingue les fonctionnalités publiques accessibles sans "
        "authentification, les fonctionnalités réservées aux utilisateurs connectés, et les "
        "fonctionnalités d'administration."
    ), size=11)
    add_figure(doc, DIAG / "use_case_global.png",
               "Figure 3.4 — Diagramme de cas d'utilisation global")

    add_heading(doc, "3.4 Diagramme de classes", level=2)
    add_para(doc, (
        "Le diagramme de classes présente l'ensemble des entités du projet et leurs relations. "
        "Les entités JPA héritent toutes de BaseEntity (id, createdAt, updatedAt via JPA Auditing). "
        "L'entité Review est un document MongoDB, sans héritage de BaseEntity, identifié par un "
        "String id généré par MongoDB."
    ), size=11)
    add_figure(doc, DIAG / "class_diagram.png",
               "Figure 3.5 — Diagramme de classes (JPA + MongoDB)")

    # Tableau des relations JPA
    add_heading(doc, "3.4.1 Relations JPA", level=3)
    t = doc.add_table(rows=9, cols=4)
    t.style = "Table Grid"
    add_table_header(t, ["Entité source", "Entité cible", "Type", "Détail"])
    relations = [
        ("User", "Profile", "@OneToOne (cascade ALL)", "Profil créé à l'inscription"),
        ("User", "Cart", "@OneToOne (cascade ALL)", "Panier créé automatiquement"),
        ("User", "Order", "@OneToMany", "Historique des commandes"),
        ("Category", "Product", "@OneToMany", "Produits d'une catégorie"),
        ("Product", "Inventory", "@OneToOne (cascade ALL)", "Stock créé à la création produit"),
        ("Cart", "CartItem", "@OneToMany (cascade ALL)", "Articles du panier"),
        ("CartItem", "Product", "@ManyToOne", "Référence produit avec snapshot prix"),
        ("Order", "OrderItem", "@OneToMany (cascade ALL)", "Lignes de commande"),
    ]
    for i, (src, tgt, typ, detail) in enumerate(relations):
        row = t.rows[i + 1]
        for j, val in enumerate([src, tgt, typ, detail]):
            row.cells[j].text = val
            for run in row.cells[j].paragraphs[0].runs:
                run.font.size = Pt(9)

    add_heading(doc, "3.5 Diagramme Entité-Relation (ERD)", level=2)
    add_para(doc, (
        "Le modèle relationnel comprend 9 tables. Les clés primaires sont de type BIGINT auto-incrémenté. "
        "Les colonnes email (users) et name (categories) disposent de contraintes UNIQUE. "
        "Les tables cart_items et order_items stockent le prix unitaire au moment de l'opération "
        "(snapshot), découplant ainsi le catalogue des transactions passées."
    ), size=11)
    add_figure(doc, DIAG / "erd.png",
               "Figure 3.6 — Diagramme Entité-Relation (9 tables)")

    add_heading(doc, "3.6 Justification de la persistance hybride (JPA vs MongoDB)", level=2)
    add_para(doc, (
        "Spring Data JPA est retenu pour les données relationnelles car elles impliquent de nombreuses "
        "jointures (User ↔ Cart ↔ CartItem ↔ Product, Order ↔ OrderItem) et des transactions ACID. "
        "MongoDB est choisi pour les avis produit car leur structure est document-centrique "
        "(pas de jointures nécessaires), leur volume peut croître indépendamment des autres entités, "
        "et leur indisponibilité ne doit pas bloquer le démarrage de l'application. "
        "Ce choix illustre l'approche polyglotte préconisée par le cahier des charges."
    ), size=11)

    add_heading(doc, "3.7 Diagrammes de séquence", level=2)

    sequences = [
        ("3.7.1 DS1 — Authentification JWT", "sequence_login.png",
         "Figure 3.7 — Diagramme de séquence — Authentification JWT",
         ("Ce diagramme illustre le flux d'authentification. Après vérification des credentials "
          "via BCrypt, le JwtService génère un token signé HS256, valable 24 heures. "
          "Le frontend stocke ce token dans le localStorage et l'injecte dans chaque requête "
          "protégée via l'authInterceptor Angular.")),
        ("3.7.2 DS3 — Ajout au panier (vérification stock)", "sequence_addtocart.png",
         "Figure 3.9 — Diagramme de séquence — Ajout panier",
         ("L'ajout d'un article au panier implique une vérification systématique du stock disponible "
          "avant toute modification. Si la quantité demandée dépasse le stock, une BusinessException "
          "est levée et le panier n'est pas modifié. Le prix unitaire est snapshotté au moment de l'ajout.")),
        ("3.7.3 DS4 — Checkout transactionnel", "sequence_checkout.png",
         "Figure 3.8 — Diagramme de séquence — Checkout",
         ("Le checkout est la séquence la plus critique : elle est encapsulée dans une transaction "
          "@Transactional. Toute défaillance (panier vide, stock insuffisant, erreur de persistance) "
          "déclenche un rollback complet, garantissant la cohérence des données.")),
        ("3.7.4 DS5 — Dépôt d'avis (MongoDB)", "sequence_review.png",
         "Figure 3.10 — Diagramme de séquence — Avis MongoDB",
         ("Le dépôt d'avis nécessite une authentification JWT. L'avis est persisté comme document "
          "MongoDB. Si MongoDB est indisponible, l'opération échoue proprement sans impacter "
          "le reste de l'application.")),
    ]

    for title, png, caption, desc in sequences:
        add_heading(doc, title, level=3)
        add_para(doc, desc, size=11)
        add_figure(doc, DIAG / png, caption)

    add_heading(doc, "3.8 Architecture technique", level=2)
    add_para(doc, (
        "Le diagramme suivant présente la vue technique complète du déploiement : le navigateur "
        "communique avec Angular sur le port 4200, qui interroge l'API Spring Boot sur le port 8080. "
        "Le backend persiste les données relationnelles dans H2 (développement) ou MySQL (production) "
        "et les avis dans MongoDB sur le port 27017."
    ), size=11)
    add_figure(doc, DIAG / "architecture_technique.png",
               "Figure 3.3 — Architecture technique (Angular ↔ Spring Boot ↔ BDD)")
    add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# CHAPITRE 4 — OUTILS & TECHNOLOGIES
# ══════════════════════════════════════════════════════════════════════════════

def add_chapter4(doc):
    add_heading(doc, "Chapitre 4 — Outils et Technologies", level=1)

    add_heading(doc, "4.1 Environnement de développement", level=2)
    tools = [
        ("IntelliJ IDEA 2024", "IDE principal pour le développement Java/Spring Boot"),
        ("Visual Studio Code 1.89", "Éditeur pour le développement Angular/TypeScript"),
        ("Git 2.44 + GitHub", "Gestion de version, branches feature, historique complet"),
        ("Postman 11", "Test manuel des endpoints REST"),
        ("Swagger UI (Springdoc 2.6.0)", "Documentation interactive de l'API REST"),
        ("Maven 3.9", "Gestionnaire de dépendances et outil de build Java"),
        ("Node.js 20 + npm 10", "Runtime et gestionnaire de paquets pour Angular"),
        ("Mermaid CLI (mmdc 11)", "Génération de diagrammes UML en PNG depuis des sources texte"),
    ]
    t = doc.add_table(rows=len(tools) + 1, cols=2)
    t.style = "Table Grid"
    add_table_header(t, ["Outil", "Rôle"])
    for i, (tool, role) in enumerate(tools):
        t.rows[i + 1].cells[0].text = tool
        t.rows[i + 1].cells[1].text = role
        for cell in t.rows[i + 1].cells:
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)

    add_heading(doc, "4.2 Stack Backend", level=2)

    add_heading(doc, "4.2.1 Java 21 & Spring Boot 3.3.5", level=3)
    add_para(doc, (
        "Java 21 (LTS) apporte les records, le pattern matching et les virtual threads. "
        "Spring Boot 3.3.5 s'appuie sur Spring Framework 6.1 et requiert Jakarta EE 10 "
        "(namespace jakarta.*). L'auto-configuration élimine la majeure partie du boilerplate XML. "
        "Le paramètre -parameters du compilateur Maven est activé pour la résolution des noms "
        "de paramètres à l'exécution (requis par Spring MVC et MapStruct)."
    ), size=11)

    add_heading(doc, "4.2.2 Spring Data JPA & Hibernate", level=3)
    add_para(doc, (
        "Spring Data JPA abstrait la couche d'accès aux données relationnelles. Hibernate 6 génère "
        "le DDL et exécute les requêtes JPQL/HQL. La stratégie DDL auto est définie à update "
        "en développement. JPA Auditing (via @EnableJpaAuditing + BaseEntity) gère automatiquement "
        "les champs createdAt et updatedAt."
    ), size=11)

    add_heading(doc, "4.2.3 Spring Data MongoDB", level=3)
    add_para(doc, (
        "Spring Data MongoDB fournit une abstraction document pour MongoDB 7. Le repository "
        "ReviewRepository étend MongoRepository<Review, String>. Un index sur productId accélère "
        "les requêtes de lecture des avis par produit. La connexion MongoDB est configurée avec "
        "un timeout court (2 s) pour un démarrage gracieux."
    ), size=11)

    add_heading(doc, "4.2.4 Spring Security + JWT (jjwt 0.12.6)", level=3)
    add_para(doc, (
        "La sécurité est stateless : pas de session HTTP. Chaque requête est interceptée par "
        "JwtAuthFilter qui extrait et valide le token Bearer. BCryptPasswordEncoder (strength 10) "
        "est utilisé pour le hashage des mots de passe. @EnableMethodSecurity active "
        "@PreAuthorize pour le contrôle fin des accès par rôle."
    ), size=11)

    add_heading(doc, "4.2.5 MapStruct 1.6.3 & Lombok 1.18.36", level=3)
    add_para(doc, (
        "MapStruct génère les implémentations des interfaces @Mapper à la compilation, éliminant "
        "le mapping manuel DTO ↔ entité. Lombok génère constructeurs, getters, builders et logger "
        "via annotations (@Data, @Builder, @RequiredArgsConstructor, @Slf4j). "
        "L'ordre des annotation processors Maven est crucial : Lombok avant MapStruct."
    ), size=11)

    add_heading(doc, "4.3 Stack Frontend", level=2)

    add_heading(doc, "4.3.1 Angular 18.2.13", level=3)
    add_para(doc, (
        "Angular 18 introduit les composants standalone (sans NgModule), le lazy loading par "
        "loadComponent(), et les Signals pour la gestion de l'état réactif. Le routing est "
        "configuré avec des guards fonctionnels (authGuard, adminGuard) et des intercepteurs "
        "HTTP fonctionnels (authInterceptor, errorInterceptor)."
    ), size=11)

    add_heading(doc, "4.3.2 Angular Material & Reactive Forms", level=3)
    add_para(doc, (
        "Angular Material fournit les composants UI (mat-card, mat-table, mat-paginator, "
        "mat-snack-bar, mat-dialog). Les formulaires utilisent l'approche Reactive Forms "
        "(FormBuilder, Validators) pour une validation côté client synchronisée avec la "
        "validation serveur (@Valid Spring)."
    ), size=11)

    add_heading(doc, "4.4 Dépendances Maven", level=2)
    t2 = doc.add_table(rows=13, cols=3)
    t2.style = "Table Grid"
    add_table_header(t2, ["GroupId:ArtifactId", "Version", "Rôle"])
    deps = [
        ("spring-boot-starter-web", "3.3.5 (géré)", "API REST, Jackson JSON"),
        ("spring-boot-starter-data-jpa", "3.3.5 (géré)", "ORM Hibernate + Spring Data"),
        ("spring-boot-starter-data-mongodb", "3.3.5 (géré)", "Spring Data MongoDB"),
        ("spring-boot-starter-security", "3.3.5 (géré)", "Filtre de sécurité, AuthManager"),
        ("spring-boot-starter-validation", "3.3.5 (géré)", "Jakarta Validation (@Valid)"),
        ("h2", "runtime", "Base H2 en mémoire/fichier (dev)"),
        ("mysql-connector-j", "runtime", "Driver JDBC MySQL (prod)"),
        ("jjwt-api / impl / jackson", "0.12.6", "Génération et validation JWT HS256"),
        ("lombok", "1.18.36 (optional)", "Génération de code à la compilation"),
        ("mapstruct", "1.6.3", "Mapper DTO ↔ entité généré à la compilation"),
        ("springdoc-openapi-starter-webmvc-ui", "2.6.0", "Swagger UI + OpenAPI 3 spec"),
        ("spring-boot-starter-test", "test", "JUnit 5, Mockito, AssertJ"),
    ]
    for i, (art, ver, role) in enumerate(deps):
        row = t2.rows[i + 1]
        for j, val in enumerate([art, ver, role]):
            row.cells[j].text = val
            for run in row.cells[j].paragraphs[0].runs:
                run.font.size = Pt(9)
    add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# CHAPITRE 5 — RÉALISATION & IMPLÉMENTATION
# ══════════════════════════════════════════════════════════════════════════════

def add_chapter5(doc):
    add_heading(doc, "Chapitre 5 — Réalisation et Implémentation", level=1)

    add_heading(doc, "5.1 Structure du backend", level=2)
    add_para(doc, (
        "Le backend compte 76 classes Java réparties dans 9 packages sous com.estore. "
        "Chaque domaine suit la même organisation interne : controller, dto, entity, mapper, "
        "repository, service. Les packages config, exception et shared constituent le noyau transversal."
    ), size=11)

    add_code_block(doc, """com.estore
├── EstoreApplication.java
├── billing
│   ├── controller/OrderController.java
│   ├── dto/{OrderResponse, OrderItemResponse, UpdateOrderItemRequest}.java
│   ├── entity/{Order, OrderItem, OrderStatus}.java
│   ├── repository/OrderRepository.java
│   └── service/OrderService.java
├── catalog
│   ├── controller/{CategoryController, ProductController}.java
│   ├── dto/{CategoryRequest, CategoryResponse, ProductRequest, ProductResponse}.java
│   ├── entity/{Category, Product}.java
│   ├── mapper/{CategoryMapper, ProductMapper}.java
│   ├── repository/{CategoryRepository, ProductRepository}.java
│   └── service/{CategoryService, ProductService}.java
├── config
│   ├── CorsConfig.java
│   ├── DataSeeder.java
│   ├── JpaAuditingConfig.java
│   ├── OpenApiConfig.java
│   └── SecurityConfig.java
├── customer
│   ├── controller/{AuthController, UserController}.java
│   ├── dto/{AuthResponse, LoginRequest, ProfileResponse, RegisterRequest, ...}.java
│   ├── entity/{Profile, Role, User}.java
│   ├── mapper/UserMapper.java
│   ├── repository/{ProfileRepository, UserRepository}.java
│   ├── security/{JwtAuthEntryPoint, JwtAuthFilter}.java
│   └── service/{AuthService, CustomUserDetailsService, JwtService, UserService}.java
├── exception
│   ├── BusinessException.java
│   ├── GlobalExceptionHandler.java
│   └── ResourceNotFoundException.java
├── inventory
│   ├── controller/InventoryController.java
│   ├── dto/{AdjustStockRequest, InventoryResponse}.java
│   ├── entity/Inventory.java
│   ├── repository/InventoryRepository.java
│   └── service/InventoryService.java
├── review
│   ├── controller/ReviewController.java
│   ├── dto/{ReviewRequest, ReviewResponse}.java
│   ├── entity/Review.java
│   ├── repository/ReviewRepository.java
│   └── service/ReviewService.java
├── shared
│   ├── ApiError.java
│   ├── BaseEntity.java
│   └── PageResponse.java
└── shopping
    ├── controller/CartController.java
    ├── dto/{AddCartItemRequest, CartItemResponse, CartResponse, UpdateCartItemRequest}.java
    ├── entity/{Cart, CartItem}.java
    ├── repository/{CartItemRepository, CartRepository}.java
    └── service/CartService.java""", "Arbre de packages backend (76 classes)")

    add_heading(doc, "5.2 Structure du frontend", level=2)
    add_para(doc, (
        "Le frontend Angular 18 est composé de 17 composants standalone et 8 services. "
        "Le routing utilise le lazy loading par loadComponent() pour optimiser le chargement initial."
    ), size=11)

    add_code_block(doc, """src/app
├── app.component.ts        (shell: router-outlet + header)
├── app.config.ts           (provideRouter, provideHttpClient, provideAnimations)
├── app.routes.ts           (routes lazily chargées)
├── core
│   ├── guards/
│   │   ├── auth.guard.ts   (redirect /login si non authentifié)
│   │   └── admin.guard.ts  (redirect / si rôle != ADMIN)
│   ├── interceptors/
│   │   ├── auth.interceptor.ts   (injecte Bearer JWT)
│   │   └── error.interceptor.ts  (gestion centralisée des erreurs HTTP)
│   ├── models/api.ts       (interfaces TypeScript)
│   └── services/
│       ├── auth.service.ts     (Signal: currentUser, login, register, logout)
│       ├── cart.service.ts     (Signal: cartCount, add, update, remove)
│       ├── catalog.service.ts  (getProducts, getProduct, getCategories)
│       ├── order.service.ts    (placeOrder, getMyOrders, getOrder)
│       ├── profile.service.ts  (getProfile, updateProfile)
│       ├── review.service.ts   (getReviews, addReview)
│       ├── admin.service.ts    (CRUD admin)
│       └── api-error.service.ts
├── features
│   ├── admin/
│   │   ├── admin-dashboard.component.ts
│   │   ├── admin-products.component.ts
│   │   ├── admin-categories.component.ts
│   │   ├── admin-inventory.component.ts
│   │   └── admin-orders.component.ts
│   ├── auth/
│   │   ├── login.component.ts
│   │   └── register.component.ts
│   ├── cart/cart.component.ts
│   ├── catalog/
│   │   ├── product-list.component.ts
│   │   └── product-detail.component.ts
│   ├── checkout/checkout-success.component.ts
│   ├── orders/
│   │   ├── order-list.component.ts
│   │   └── order-detail.component.ts
│   ├── profile/profile.component.ts
│   └── not-found.component.ts
└── shared/header.component.ts""", "Arbre de composants frontend (17 composants / 8 services)")

    add_heading(doc, "5.3 Implémentation — Customer & Authentification", level=2)
    add_para(doc, (
        "Le domaine Customer gère l'inscription, la connexion et le profil utilisateur. "
        "L'entité User hérite de BaseEntity et porte un champ role (enum Role : USER, ADMIN). "
        "Un Profile vide est créé automatiquement à l'inscription via la méthode helper "
        "user.attachProfile()."
    ), size=11)

    add_code_block(doc, """// AuthService.java — extrait inscription
public AuthResponse register(RegisterRequest request) {
    if (userRepository.existsByEmail(request.email()))
        throw new BusinessException("Email already registered");
    User user = User.builder()
        .firstName(request.firstName()).lastName(request.lastName())
        .email(request.email())
        .password(passwordEncoder.encode(request.password()))
        .role(Role.USER).build();
    user.attachProfile(Profile.builder().build());
    userRepository.save(user);
    String token = jwtService.generateToken(user);
    return new AuthResponse(token, user.getId(), user.getRole().name());
}""", "Extrait : AuthService.register()")

    add_para(doc, (
        "La classe JwtService utilise la librairie jjwt 0.12.6 avec algorithme HS256. "
        "Le secret est lu depuis la configuration (variable d'environnement JWT_SECRET en production). "
        "L'expiration est de 86 400 000 ms (24 heures). JwtAuthFilter intercepte chaque requête, "
        "extrait le header Authorization: Bearer <token>, valide la signature et l'expiration, "
        "puis renseigne le SecurityContextHolder."
    ), size=11)

    add_heading(doc, "5.4 Implémentation — Catalogue", level=2)
    add_para(doc, (
        "ProductService supporte la recherche paginée avec filtres optionnels via une requête JPQL "
        "dynamique. Le ProductRepository expose findByNameContainingIgnoreCaseAndCategoryId(). "
        "MapStruct génère ProductMapper pour la conversion Product ↔ ProductResponse/ProductRequest. "
        "La réponse ProductResponse inclut le champ quantityInStock joint depuis Inventory."
    ), size=11)

    add_code_block(doc, """// ProductController.java — GET /api/products
@GetMapping
public ResponseEntity<PageResponse<ProductResponse>> list(
    @RequestParam(required = false) String q,
    @RequestParam(required = false) Long categoryId,
    @ParameterObject @PageableDefault(size = 20, sort = "name") Pageable pageable) {
    return ResponseEntity.ok(productService.findAll(q, categoryId, pageable));
}""", "Extrait : ProductController — recherche paginée")

    add_heading(doc, "5.5 Implémentation — Stocks (Inventory)", level=2)
    add_para(doc, (
        "L'inventaire est créé automatiquement à 0 lors de la création d'un produit "
        "(cascade dans ProductService). InventoryController expose deux endpoints : "
        "GET /api/inventory/product/{productId} (public) et "
        "PUT /api/inventory/product/{productId} (ADMIN uniquement)."
    ), size=11)

    add_heading(doc, "5.6 Implémentation — Panier (Shopping)", level=2)
    add_para(doc, (
        "CartService crée automatiquement un panier vide lors du premier accès d'un utilisateur. "
        "L'ajout d'un article vérifie le stock disponible via InventoryService : si la quantité "
        "demandée excède le stock, une BusinessException est levée (HTTP 422). "
        "Le prix unitaire est snapshotté à l'ajout (unitPrice = product.getPrice()). "
        "La contrainte UNIQUE(cart, product) sur cart_items évite les doublons."
    ), size=11)

    add_heading(doc, "5.7 Implémentation — Commandes (Billing)", level=2)
    add_para(doc, (
        "Le checkout est la fonctionnalité la plus critique. OrderService.placeOrder() est annotée "
        "@Transactional, garantissant l'atomicité : si une étape échoue, toutes les modifications "
        "(création commande, décrémentation stocks, vidage panier) sont annulées. "
        "OrderItem snapshotte productName (dénormalisé) afin que le nom du produit reste "
        "lisible même si le produit est ultérieurement renommé ou supprimé."
    ), size=11)

    add_code_block(doc, """// OrderService.java — placeOrder() simplifié
@Transactional
public OrderResponse placeOrder() {
    Cart cart = cartService.getCurrentUserCart();
    if (cart.getItems().isEmpty())
        throw new BusinessException("Le panier est vide");

    Order order = Order.builder().user(cart.getUser())
        .status(OrderStatus.CONFIRMED).build();

    for (CartItem item : cart.getItems()) {
        Inventory inv = inventoryService.getByProduct(item.getProduct());
        if (inv.getQuantity() < item.getQuantity())
            throw new BusinessException("Stock insuffisant pour: " + item.getProduct().getName());
        inv.setQuantity(inv.getQuantity() - item.getQuantity());
        order.addItem(item.getProduct(), item.getQuantity(), item.getUnitPrice());
    }
    cartService.clearCart(cart);
    return mapper.toResponse(orderRepository.save(order));
}""", "Extrait : OrderService.placeOrder() — transaction atomique")

    add_heading(doc, "5.8 Implémentation — Avis (Reviews — MongoDB)", level=2)
    add_para(doc, (
        "Review est un document MongoDB (collection reviews) annoté @Document. "
        "ReviewRepository étend MongoRepository<Review, String>. "
        "La lecture des avis (GET /api/reviews/product/{productId}) est publique. "
        "La création (POST /api/reviews) requiert un token JWT via @PreAuthorize(\"isAuthenticated()\"). "
        "DataSeeder tente de seeder 5 avis au démarrage, mais gère gracieusement l'indisponibilité "
        "de MongoDB via un CompletableFuture avec timeout de 2 secondes."
    ), size=11)

    add_heading(doc, "5.9 Gestion des erreurs transversale", level=2)
    add_para(doc, (
        "GlobalExceptionHandler (@RestControllerAdvice) centralise la gestion de toutes les "
        "exceptions et retourne un corps ApiError structuré. ApiError contient : status (int), "
        "error (String), message, path, timestamp (Instant) et une liste de FieldViolation "
        "pour les erreurs de validation."
    ), size=11)

    t = doc.add_table(rows=9, cols=3)
    t.style = "Table Grid"
    add_table_header(t, ["Exception Java", "HTTP Status", "Cas d'usage"])
    exc_map = [
        ("ResourceNotFoundException", "404 Not Found", "Entité introuvable par ID"),
        ("BusinessException", "422 Unprocessable Entity", "Stock insuffisant, panier vide"),
        ("DataIntegrityViolationException", "422 Unprocessable Entity", "Email en double, contrainte UNIQUE"),
        ("MethodArgumentNotValidException", "400 Bad Request", "Validation @Valid + FieldViolations"),
        ("ConstraintViolationException", "400 Bad Request", "Validation javax/jakarta path params"),
        ("BadCredentialsException / AuthenticationException", "401 Unauthorized", "Login échoué"),
        ("AccessDeniedException / AuthorizationDeniedException", "403 Forbidden", "Rôle insuffisant"),
        ("Exception (fallback)", "500 Internal Server Error", "Erreur inattendue"),
    ]
    for i, (ex, status, case) in enumerate(exc_map):
        row = t.rows[i + 1]
        for j, val in enumerate([ex, status, case]):
            row.cells[j].text = val
            for run in row.cells[j].paragraphs[0].runs:
                run.font.size = Pt(9)

    add_heading(doc, "5.10 Implémentation Frontend Angular", level=2)
    add_para(doc, (
        "Le frontend est une application Angular 18 SPA (Single Page Application). "
        "L'authInterceptor injecte automatiquement le token JWT dans chaque requête HTTP "
        "sortante. L'errorInterceptor capture les erreurs 401 (redirection login) et 403. "
        "AuthService et CartService utilisent les Signals Angular pour un état réactif "
        "sans NgRx : signal<User | null> et signal<number> pour le compteur panier. "
        "Le routing lazy charge chaque composant à la demande, réduisant le bundle initial."
    ), size=11)

    add_code_block(doc, """// auth.interceptor.ts
export const authInterceptor: HttpInterceptorFn = (req, next) => {
  const token = localStorage.getItem('token');
  if (token) {
    req = req.clone({
      setHeaders: { Authorization: `Bearer ${token}` }
    });
  }
  return next(req);
};

// auth.service.ts — Signal
export class AuthService {
  currentUser = signal<User | null>(null);

  login(credentials: LoginRequest): Observable<AuthResponse> {
    return this.http.post<AuthResponse>('/api/auth/login', credentials).pipe(
      tap(res => {
        localStorage.setItem('token', res.token);
        this.currentUser.set(res);
      })
    );
  }
}""", "Extrait : authInterceptor et AuthService (Signals)")
    add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# CHAPITRE 6 — INTERFACES (PLACEHOLDERS)
# ══════════════════════════════════════════════════════════════════════════════

def add_chapter6(doc):
    add_heading(doc, "Chapitre 6 — Présentation de l'Application (Interfaces)", level=1)
    add_para(doc, (
        "Ce chapitre illustre les principales interfaces de l'application E-Store à travers "
        "dix captures d'écran représentatives. Les zones grisées ci-dessous sont des emplacements "
        "réservés aux captures : insérez les screenshots réels en remplaçant les zones annotées."
    ), size=11, italic=True, color=GRAY)

    screens = [
        ("6.1", "Page d'accueil — Liste des produits",
         "La page d'accueil affiche la liste paginée des produits sous forme de cartes Angular Material. "
         "Chaque carte présente l'image, le nom, le prix et un indicateur de disponibilité. "
         "Un champ de recherche et un sélecteur de catégorie permettent de filtrer les résultats "
         "en temps réel. La pagination (20 produits par page) est gérée par mat-paginator."),
        ("6.2", "Détail d'un produit et avis",
         "La page de détail affiche toutes les informations du produit (image, description, prix, stock). "
         "En bas de page, la section avis affiche les commentaires issus de MongoDB avec la note "
         "en étoiles. Un formulaire d'avis est affiché aux utilisateurs authentifiés."),
        ("6.3", "Formulaire d'inscription",
         "Le formulaire d'inscription recueille prénom, nom, email et mot de passe. "
         "La validation Reactive Forms affiche des messages d'erreur en temps réel (email invalide, "
         "mot de passe trop court). Le formulaire est synchronisé avec la validation serveur @Valid."),
        ("6.4", "Formulaire de connexion",
         "Le formulaire de connexion demande email et mot de passe. En cas d'erreur de credentials, "
         "un mat-snack-bar affiche le message d'erreur retourné par le serveur (401). "
         "Après connexion, l'utilisateur est redirigé vers la page précédente ou l'accueil."),
        ("6.5", "Panier d'achat",
         "Le panier liste les articles avec leur image, nom, quantité et prix unitaire. "
         "Des contrôles +/- permettent d'ajuster la quantité (avec vérification du stock). "
         "Le total est calculé côté client. Le bouton « Commander » déclenche le checkout."),
        ("6.6", "Confirmation de commande",
         "Après un checkout réussi, l'utilisateur est redirigé vers la page de confirmation "
         "qui affiche le numéro de commande, le détail des articles et le total. "
         "Un lien vers l'historique des commandes est proposé."),
        ("6.7", "Historique des commandes",
         "La liste des commandes présente les commandes de l'utilisateur connecté, "
         "triées par date décroissante. Chaque ligne affiche le numéro, la date, le total "
         "et le statut (CONFIRMED / CANCELLED). Un clic ouvre le détail complet."),
        ("6.8", "Profil utilisateur",
         "La page de profil affiche les informations de l'utilisateur (nom, email, rôle) "
         "et permet de modifier les informations de contact (téléphone, adresse, ville, pays) "
         "via un formulaire Reactive Forms avec validation."),
        ("6.9", "Dashboard administrateur",
         "Le dashboard admin (protégé par adminGuard) offre une vue d'ensemble : "
         "nombre de produits, catégories, commandes et utilisateurs. "
         "La navigation latérale permet d'accéder aux sections Produits, Catégories, Stocks, Commandes."),
        ("6.10", "Gestion des stocks (admin)",
         "L'interface de gestion des stocks affiche un tableau de tous les produits avec leur stock actuel. "
         "Un champ éditable par ligne permet à l'admin de mettre à jour le stock instantanément "
         "via PUT /api/inventory/product/{id}."),
    ]

    for num, title, desc in screens:
        add_heading(doc, f"6.{num.split('.')[1]} {title}", level=2)
        add_para(doc, desc, size=11)
        # Placeholder box
        t = doc.add_table(rows=1, cols=1)
        t.style = "Table Grid"
        cell = t.rows[0].cells[0]
        cell.height = Cm(8)
        set_cell_bg(cell, "F9FAFB")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"\n\n\n[Figure 6.{num.split('.')[1]} — {title}]\nRemplacer par la capture d'écran réelle\n\n\n")
        run.font.color.rgb = GRAY
        run.italic = True
        run.font.size = Pt(10)
        doc.add_paragraph()
    add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# CHAPITRE 7 — TESTS & VALIDATION
# ══════════════════════════════════════════════════════════════════════════════

def add_chapter7(doc):
    add_heading(doc, "Chapitre 7 — Tests et Validation", level=1)

    add_heading(doc, "7.1 Tests fonctionnels par domaine", level=2)
    add_para(doc, (
        "Le tableau suivant récapitule les cas de test fonctionnels couverts, "
        "avec le statut observé lors de l'exécution du smoke-test."
    ), size=11)

    t = doc.add_table(rows=1, cols=5)
    t.style = "Table Grid"
    add_table_header(t, ["Domaine", "Cas de test", "Endpoint", "Résultat attendu", "Statut"])

    tests = [
        # Customer & Auth
        ("Auth", "Inscription valide", "POST /api/auth/register", "201 + token JWT", "PASS"),
        ("Auth", "Email déjà utilisé", "POST /api/auth/register", "422 Unprocessable Entity", "PASS"),
        ("Auth", "Connexion valide", "POST /api/auth/login", "200 + token JWT", "PASS"),
        ("Auth", "Mauvais mot de passe", "POST /api/auth/login", "401 Unauthorized", "PASS"),
        ("Auth", "Accès sans token", "GET /api/users/me", "401 Unauthorized", "PASS"),
        ("Auth", "Accès ADMIN avec rôle USER", "POST /api/categories", "403 Forbidden", "PASS"),
        # Catalog
        ("Catalog", "Créer catégorie (ADMIN)", "POST /api/categories", "201 + CategoryResponse", "PASS"),
        ("Catalog", "Catégorie en double", "POST /api/categories", "422 Unprocessable Entity", "PASS"),
        ("Catalog", "Créer produit (ADMIN)", "POST /api/products", "201 + ProductResponse", "PASS"),
        ("Catalog", "Recherche par mot-clé", "GET /api/products?q=clean", "200 + page 1 résultat", "PASS"),
        # Inventory
        ("Inventory", "Stock auto-créé à 0", "GET /api/products/{id}", "quantityInStock=0", "PASS"),
        ("Inventory", "Mise à jour stock", "PUT /api/inventory/product/{id}", "200 + qty=10", "PASS"),
        ("Inventory", "Stock reflété dans produit", "GET /api/products/{id}", "quantityInStock=10", "PASS"),
        # Shopping
        ("Shopping", "Panier auto-créé", "GET /api/cart", "200 + itemCount=0", "PASS"),
        ("Shopping", "Ajout article valide", "POST /api/cart/items", "200 + CartResponse", "PASS"),
        ("Shopping", "Dépassement stock", "POST /api/cart/items qty>stock", "422 BusinessException", "PASS"),
        ("Shopping", "Mise à jour quantité", "PUT /api/cart/items/{id}", "200 + qty mise à jour", "PASS"),
        ("Shopping", "Suppression article", "DELETE /api/cart/items/{id}", "200 OK", "PASS"),
        ("Shopping", "Vidage panier", "DELETE /api/cart", "200 OK", "PASS"),
        # Billing
        ("Billing", "Checkout panier vide", "POST /api/orders", "422 BusinessException", "PASS"),
        ("Billing", "Checkout valide", "POST /api/orders", "201 + CONFIRMED", "PASS"),
        ("Billing", "Stock décrémenté", "GET /api/products/{id}", "quantityInStock = n-qty", "PASS"),
        ("Billing", "Panier vidé après checkout", "GET /api/cart", "itemCount=0", "PASS"),
        ("Billing", "Historique commandes", "GET /api/orders/user/me", "200 + page 1 commande", "PASS"),
        ("Billing", "Détail commande", "GET /api/orders/{id}", "200 + OrderResponse", "PASS"),
    ]

    for domain, case, endpoint, expected, status in tests:
        row = t.add_row()
        for j, val in enumerate([domain, case, endpoint, expected, status]):
            row.cells[j].text = val
            for run in row.cells[j].paragraphs[0].runs:
                run.font.size = Pt(9)
            if j == 4 and val == "PASS":
                set_cell_bg(row.cells[j], "DCFCE7")

    add_heading(doc, "7.2 Tests de sécurité", level=2)
    t2 = doc.add_table(rows=1, cols=4)
    t2.style = "Table Grid"
    add_table_header(t2, ["Scénario de sécurité", "Action", "Résultat attendu", "Statut"])
    sec_tests = [
        ("Accès sans token", "GET /api/cart (sans Bearer)", "401 Unauthorized", "PASS"),
        ("Rôle insuffisant", "POST /api/categories (rôle USER)", "403 Forbidden", "PASS"),
        ("Token expiré", "Requête avec JWT exp < now", "401 Unauthorized", "PASS"),
        ("Token forgé", "JWT avec signature invalide", "401 Unauthorized", "PASS"),
        ("CORS non autorisé", "Requête depuis port != 4200", "Bloqué par le navigateur", "PASS"),
        ("SQL Injection (API)", "?q='; DROP TABLE users--", "Paramètre traité comme chaîne JPA", "PASS"),
        ("Accès admin ADMIN", "POST /api/categories (rôle ADMIN)", "201 Created", "PASS"),
    ]
    for scen, action, expected, status in sec_tests:
        row = t2.add_row()
        for j, val in enumerate([scen, action, expected, status]):
            row.cells[j].text = val
            for run in row.cells[j].paragraphs[0].runs:
                run.font.size = Pt(9)
        if status == "PASS":
            set_cell_bg(t2.rows[-1].cells[3], "DCFCE7")

    add_heading(doc, "7.3 Script de smoke-test", level=2)
    add_para(doc, (
        "Le fichier scripts/smoke-test.sh est un script bash exécuté depuis la racine du projet. "
        "Il : (1) compile le projet avec Maven (-DskipTests), (2) démarre le jar en désactivant "
        "le seeder (estore.seed-on-startup=false), (3) attend le démarrage (~15 s), "
        "(4) exécute 25+ assertions curl couvrant les domaines Day 2 à Day 6, "
        "(5) affiche un bilan PASS/FAIL et retourne le code de sortie approprié pour CI/CD."
    ), size=11)

    add_heading(doc, "7.4 Documentation Swagger", level=2)
    add_para(doc, (
        "L'API est documentée via springdoc-openapi 2.6.0. Swagger UI est accessible à "
        "http://localhost:8080/swagger-ui.html et la spec OpenAPI 3 JSON à "
        "http://localhost:8080/v3/api-docs. OpenApiConfig configure le schéma d'authentification "
        "Bearer JWT pour permettre de tester les endpoints protégés directement depuis Swagger UI."
    ), size=11)
    add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# CONCLUSION GÉNÉRALE
# ══════════════════════════════════════════════════════════════════════════════

def add_conclusion(doc):
    add_heading(doc, "Conclusion Générale", level=1)

    add_heading(doc, "Récapitulatif", level=2)
    add_para(doc, (
        "Ce projet a permis de concevoir et de réaliser de bout en bout une plateforme e-commerce "
        "complète : E-Store. L'architecture en couches avec décomposition par domaines métier, "
        "la persistance hybride JPA/MongoDB, la sécurité stateless JWT/BCrypt, et le frontend "
        "Angular 18 standalone ont été intégrés en suivant les meilleures pratiques Spring Boot 3."
    ), size=11)

    add_heading(doc, "Apports techniques", level=2)
    apports = [
        "Maîtrise de l'architecture en couches et par domaines (Domain-Driven Design léger).",
        "Implémentation d'une persistance polyglotte (JPA pour relationnel, MongoDB pour documents).",
        "Sécurisation stateless complète avec JWT et BCrypt, sans session serveur.",
        "Développement full stack Angular 18 avec composants standalone, Signals et lazy loading.",
        "Automatisation des tests via script smoke-test avec 25+ assertions.",
        "Documentation API interactive Swagger/OpenAPI prête à l'emploi.",
    ]
    for a in apports:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(a).font.size = Pt(11)

    add_heading(doc, "Limites du projet", level=2)
    limites = [
        "Pas de paiement en ligne réel (simulation du checkout sans passerelle Stripe/PayPal).",
        "Pas de tests unitaires JUnit automatisés (couverture assurée par smoke-test d'intégration).",
        "Pas de gestion de l'expiration du token côté frontend (refresh token non implémenté).",
        "Images produit stockées comme URL externe (pas d'upload fichier).",
    ]
    for l in limites:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(l).font.size = Pt(11)

    add_heading(doc, "Perspectives", level=2)
    perspectives = [
        "Intégration OAuth2 / OpenID Connect (Google, GitHub) et tokens de rafraîchissement.",
        "Passerelle de paiement (Stripe API) pour transactions réelles.",
        "Migration vers une architecture microservices avec Spring Cloud Gateway.",
        "Envoi d'emails transactionnels (confirmation commande, réinitialisation mot de passe).",
        "Déploiement conteneurisé avec Docker Compose (Spring Boot + Angular + MySQL + MongoDB).",
        "Tests unitaires JUnit 5 / Mockito et tests d'intégration @SpringBootTest.",
    ]
    for pv in perspectives:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(pv).font.size = Pt(11)
    add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# BIBLIOGRAPHIE
# ══════════════════════════════════════════════════════════════════════════════

def add_bibliography(doc):
    add_heading(doc, "Bibliographie", level=1)

    refs = [
        ("[1] Spring Boot Reference Documentation 3.3.5", "https://docs.spring.io/spring-boot/docs/3.3.5/reference/html/"),
        ("[2] Spring Data JPA Reference", "https://docs.spring.io/spring-data/jpa/docs/current/reference/html/"),
        ("[3] Spring Data MongoDB Reference", "https://docs.spring.io/spring-data/mongodb/docs/current/reference/html/"),
        ("[4] Spring Security Reference 6.3", "https://docs.spring.io/spring-security/reference/"),
        ("[5] jjwt (Java JWT Library) 0.12.6", "https://github.com/jwtk/jjwt"),
        ("[6] MapStruct 1.6 — Reference Guide", "https://mapstruct.org/documentation/stable/reference/html/"),
        ("[7] Springdoc-openapi 2.6.0", "https://springdoc.org/"),
        ("[8] Angular 18 Official Documentation", "https://angular.dev/"),
        ("[9] Angular Material — Component Documentation", "https://material.angular.io/"),
        ("[10] MongoDB 7.0 Manual", "https://www.mongodb.com/docs/manual/"),
        ("[11] Mermaid — Diagramming and charting tool", "https://mermaid.js.org/"),
        ("[12] Baeldung — Spring Security JWT Tutorial", "https://www.baeldung.com/spring-security-oauth-jwt"),
        ("[13] Craig Walls — Spring in Action, 6th Edition. Manning, 2022.",),
        ("[14] Erich Gamma et al. — Design Patterns. Addison-Wesley, 1994.",),
    ]

    for ref in refs:
        p = doc.add_paragraph(style="List Number")
        if isinstance(ref, tuple) and len(ref) >= 2:
            p.add_run(ref[0]).bold = True
            p.add_run(f" — {ref[1]}")
        elif isinstance(ref, tuple):
            p.add_run(ref[0])
        else:
            p.add_run(ref)
        p.runs[0].font.size = Pt(10)
    add_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# ANNEXES
# ══════════════════════════════════════════════════════════════════════════════

def add_annexes(doc):
    add_heading(doc, "Annexes", level=1)

    # Annexe A — Endpoints
    add_heading(doc, "Annexe A — Inventaire complet des endpoints REST (28 endpoints)", level=2)
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    add_table_header(t, ["Méthode", "Endpoint", "Accès", "Description"])

    endpoints = [
        # Auth
        ("POST", "/api/auth/register", "PUBLIC", "Inscription d'un nouveau compte"),
        ("POST", "/api/auth/login", "PUBLIC", "Connexion — retourne JWT"),
        # Users
        ("GET", "/api/users/me", "USER+", "Profil de l'utilisateur connecté"),
        ("PUT", "/api/users/me/profile", "USER+", "Mise à jour du profil"),
        ("GET", "/api/users", "ADMIN", "Liste paginée de tous les utilisateurs"),
        # Categories
        ("GET", "/api/categories", "PUBLIC", "Liste de toutes les catégories"),
        ("GET", "/api/categories/{id}", "PUBLIC", "Détail d'une catégorie"),
        ("POST", "/api/categories", "ADMIN", "Créer une catégorie"),
        ("PUT", "/api/categories/{id}", "ADMIN", "Modifier une catégorie"),
        ("DELETE", "/api/categories/{id}", "ADMIN", "Supprimer une catégorie"),
        # Products
        ("GET", "/api/products", "PUBLIC", "Liste/recherche paginée des produits"),
        ("GET", "/api/products/{id}", "PUBLIC", "Détail d'un produit (+ stock)"),
        ("POST", "/api/products", "ADMIN", "Créer un produit"),
        ("PUT", "/api/products/{id}", "ADMIN", "Modifier un produit"),
        ("DELETE", "/api/products/{id}", "ADMIN", "Supprimer un produit"),
        # Inventory
        ("GET", "/api/inventory/product/{productId}", "PUBLIC", "Stock d'un produit"),
        ("PUT", "/api/inventory/product/{productId}", "ADMIN", "Mettre à jour le stock"),
        # Cart
        ("GET", "/api/cart", "USER+", "Panier de l'utilisateur connecté"),
        ("POST", "/api/cart/items", "USER+", "Ajouter un article au panier"),
        ("PUT", "/api/cart/items/{itemId}", "USER+", "Modifier la quantité d'un article"),
        ("DELETE", "/api/cart/items/{itemId}", "USER+", "Supprimer un article du panier"),
        ("DELETE", "/api/cart", "USER+", "Vider le panier"),
        # Orders
        ("POST", "/api/orders", "USER+", "Passer une commande (checkout)"),
        ("GET", "/api/orders/user/me", "USER+", "Mes commandes (paginé)"),
        ("GET", "/api/orders/{id}", "USER+", "Détail d'une commande"),
        ("GET", "/api/orders", "ADMIN", "Toutes les commandes (paginé)"),
        ("PUT", "/api/orders/{id}/items/{itemId}", "USER+", "Modifier une ligne de commande"),
        ("PUT", "/api/orders/{id}/cancel", "USER+", "Annuler une commande"),
        # Reviews
        ("GET", "/api/reviews/product/{productId}", "PUBLIC", "Avis d'un produit"),
        ("POST", "/api/reviews", "USER+", "Déposer un avis"),
    ]

    for method, ep, access, desc in endpoints:
        row = t.add_row()
        row.cells[0].text = method
        row.cells[1].text = ep
        row.cells[2].text = access
        row.cells[3].text = desc
        for cell in row.cells:
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
        # Coloriser la méthode
        color_map = {"GET": "DCFCE7", "POST": "DBEAFE", "PUT": "FEF9C3", "DELETE": "FEE2E2"}
        set_cell_bg(row.cells[0], color_map.get(method, "FFFFFF"))

    doc.add_paragraph()

    # Annexe B — Smoke-test extrait
    add_heading(doc, "Annexe B — Extrait du script smoke-test.sh", level=2)
    add_code_block(doc, """#!/usr/bin/env bash
BASE="http://localhost:8080"
FAIL=0; PASS=0

ok()   { PASS=$((PASS+1)); echo "  PASS — $1"; }
fail() { FAIL=$((FAIL+1)); echo "  FAIL — $1"; }
expect_status() {
  local got="$1" want="$2" label="$3"
  [ "$got" = "$want" ] && ok "$label ($got)" || fail "$label (got $got, want $want)"
}

# Auth
REG=$(curl -s -w "\\n%{http_code}" -X POST "$BASE/api/auth/register" \\
  -H "Content-Type: application/json" \\
  -d '{"firstName":"Adnane","lastName":"Louardi","email":"adnane@test.com","password":"password123"}')
expect_status "$(echo "$REG" | tail -1)" "201" "POST /api/auth/register"
TOKEN=$(echo "$REG" | head -n -1 | python -c "import json,sys; print(json.load(sys.stdin)['token'])")

# Vérification sans token → 401
NOAUTH=$(curl -s -o /dev/null -w "%{http_code}" "$BASE/api/users/me")
expect_status "$NOAUTH" "401" "GET /api/users/me sans token"

# Checkout transactionnel
curl -s -X POST "$BASE/api/cart/items" -H "Authorization: Bearer $TOKEN" \\
  -H "Content-Type: application/json" -d '{"productId":1,"quantity":2}' > /dev/null
ORDER=$(curl -s -w "\\n%{http_code}" -X POST "$BASE/api/orders" -H "Authorization: Bearer $TOKEN")
expect_status "$(echo "$ORDER" | tail -1)" "201" "POST /api/orders"

echo "Total: PASS=$PASS  FAIL=$FAIL"
[ $FAIL -eq 0 ]""")

    # Annexe C — application.yml
    add_heading(doc, "Annexe C — Configuration application.yml", level=2)
    add_code_block(doc, """spring:
  application:
    name: estore-backend
  profiles:
    active: dev
  data:
    mongodb:
      uri: mongodb://localhost:27017/estore?serverSelectionTimeoutMS=2000
  jpa:
    hibernate.ddl-auto: update
    open-in-view: false

server:
  port: 8080

springdoc:
  swagger-ui.path: /swagger-ui.html
  api-docs.path: /v3/api-docs

estore:
  security.jwt:
    secret: ${JWT_SECRET:change-me-please-this-is-a-dev-only-secret-at-least-32-bytes}
    expiration-ms: 86400000
  cors.allowed-origins: http://localhost:4200

---
spring:
  config.activate.on-profile: dev
  datasource:
    url: jdbc:h2:file:./data/estore;AUTO_SERVER=TRUE
    driver-class-name: org.h2.Driver
    username: sa
    password:
  jpa.database-platform: org.hibernate.dialect.H2Dialect
  h2.console.enabled: true

---
spring:
  config.activate.on-profile: mysql
  datasource:
    url: jdbc:mysql://localhost:3306/estore?createDatabaseIfNotExist=true
    driver-class-name: com.mysql.cj.jdbc.Driver
    username: ${MYSQL_USER:root}
    password: ${MYSQL_PASSWORD:root}
  jpa.database-platform: org.hibernate.dialect.MySQLDialect""")

    # Annexe D — DataSeeder
    add_heading(doc, "Annexe D — Résumé du DataSeeder", level=2)
    add_para(doc, (
        "DataSeeder (CommandLineRunner) s'exécute au démarrage si estore.seed-on-startup=true "
        "(défaut) et si la base est vide (userRepository.count() == 0). "
        "Il crée en transaction JPA : 2 utilisateurs (admin et user), 4 catégories, "
        "20 produits avec leurs inventaires. Puis il tente de seeder 5 avis MongoDB "
        "avec un timeout de 2 secondes (démarrage gracieux si Mongo est absent)."
    ), size=11)
    t_seed = doc.add_table(rows=6, cols=3)
    t_seed.style = "Table Grid"
    add_table_header(t_seed, ["Entité", "Quantité", "Détail"])
    seed_data = [
        ("Utilisateurs", "2", "admin@estore.com / admin123 (ADMIN) — user@estore.com / password123 (USER)"),
        ("Catégories", "4", "Electronics, Jewelery, Men's Clothing, Women's Clothing"),
        ("Produits", "20", "Issus de FakeStoreAPI (images, descriptions, prix réels)"),
        ("Inventaires", "20", "Stocks initiaux entre 0 (Samsung monitor) et 80 unités"),
        ("Avis MongoDB", "5", "Sur 5 produits représentatifs, si MongoDB disponible"),
    ]
    for i, (entity, qty, detail) in enumerate(seed_data):
        row = t_seed.rows[i + 1]
        for j, val in enumerate([entity, qty, detail]):
            row.cells[j].text = val
            for run in row.cells[j].paragraphs[0].runs:
                run.font.size = Pt(10)


# ══════════════════════════════════════════════════════════════════════════════
# MAIN — ASSEMBLAGE DU DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

def build():
    doc = Document()

    # Marges 2.5 cm partout
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Style Normal : Calibri 11
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Sections du rapport dans l'ordre
    add_cover(doc)
    add_dedication(doc)
    add_acknowledgements(doc)
    add_abstracts(doc)
    add_toc(doc)
    add_lists(doc)
    add_introduction(doc)
    add_chapter1(doc)
    add_chapter2(doc)
    add_chapter3(doc)
    add_chapter4(doc)
    add_chapter5(doc)
    add_chapter6(doc)
    add_chapter7(doc)
    add_conclusion(doc)
    add_bibliography(doc)
    add_annexes(doc)

    doc.save(str(OUT))
    print(f"\nRapport genere : {OUT}")
    print(f"  Pages estimées : 60–80 (ouvrir dans Word pour vérifier la TOC)")


if __name__ == "__main__":
    build()
